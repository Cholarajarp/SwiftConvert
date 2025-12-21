"""Conversion helpers extracted from app.py to improve cohesion and reduce complexity.
Keep functions pure and raise on errors; logging is local to this module.
"""
from pathlib import Path
import logging
from PIL import Image
import pandas as pd
from pdf2docx import Converter

logger = logging.getLogger(__name__)
# Allow importing heavy libs inside functions to avoid top-level import costs
# pylint: disable=import-outside-toplevel,broad-except


def handle_conversion_error(error, operation):
    """Log a conversion error and re-raise the original exception."""
    logger.error("%s error: %s", operation, error)
    # Re-raise the original exception to preserve traceback
    raise error


# PDF to DOCX conversion
def convert_pdf_to_docx(input_path: Path, output_path: Path):
    """Convert a PDF file to DOCX format using pdf2docx."""
    try:
        cv = Converter(str(input_path))
        cv.convert(str(output_path))
        cv.close()
        logger.info("PDF to DOCX conversion successful: %s", output_path)
        return output_path
    except Exception as e:
        handle_conversion_error(e, "PDF to DOCX")
        raise


# DOCX to PDF conversion (reportlab rendering helper kept lightweight)
def _render_paragraphs_to_pdf(paragraphs, output_path):
    """Render paragraphs into a PDF using ReportLab's high-level platypus API.

    This reduces low-level layout code and improves maintainability.
    """
    from reportlab.lib.pagesizes import letter as PAGE_LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    doc = SimpleDocTemplate(str(output_path), pagesize=PAGE_LETTER)
    styles = getSampleStyleSheet()
    story = []

    for paragraph in paragraphs:
        text = paragraph.strip()
        if not text:
            continue
        story.append(Paragraph(text.replace('\n', '<br/>'), styles['Normal']))
        story.append(Spacer(1, 6))

    doc.build(story)
    logger.info("Rendered paragraphs to PDF: %s", output_path)


def convert_docx_to_pdf(input_path: Path, output_path: Path):
    """Convert a DOCX file to PDF using simple paragraph rendering."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(input_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        _render_paragraphs_to_pdf(paragraphs, output_path)
        logger.info("DOCX to PDF conversion successful: %s", output_path)
        return output_path
    except Exception as e:
        handle_conversion_error(e, "DOCX to PDF")
        raise


# Image to PDF conversion
def convert_image_to_pdf(input_path: Path, output_path: Path):
    """Convert an image file into a single-page PDF."""
    try:
        with Image.open(str(input_path)) as img:
            if img.mode in ('RGBA', 'LA', 'P'):
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                if img.mode in ('RGBA', 'LA'):
                    rgb_img.paste(img, mask=img.split()
                                  [-1] if len(img.split()) > 3 else None)
                img = rgb_img

            img.save(str(output_path), 'PDF', resolution=100.0, quality=95)

        logger.info("Image to PDF conversion successful: %s", output_path)
        return output_path
    except Exception as e:
        handle_conversion_error(e, "Image to PDF")
        raise


# Image to Image conversion
def convert_image_to_image(
        input_path: Path,
        output_path: Path,
        target_format: str):
    """Convert an image from one format to another (e.g. PNG -> JPEG)."""
    try:
        with Image.open(str(input_path)) as img:
            if target_format.lower() in ['jpg', 'jpeg'] and img.mode == 'RGBA':
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(
                    img, mask=img.split()[3] if len(
                        img.split()) == 4 else None)
                rgb_img.save(str(output_path), 'JPEG', quality=95)
            else:
                img.save(
                    str(output_path),
                    format=target_format.upper(),
                    quality=95)
        logger.info("Image conversion successful: %s", output_path)
        return output_path
    except Exception as e:
        handle_conversion_error(e, "Image conversion")
        raise


# CSV <-> XLSX
def convert_csv_to_xlsx(input_path: Path, output_path: Path):
    """Convert a CSV file to XLSX using pandas + openpyxl."""
    try:
        df = pd.read_csv(str(input_path))
        df.to_excel(str(output_path), index=False, engine='openpyxl')
        logger.info("CSV to XLSX conversion successful: %s", output_path)
        return output_path
    except Exception as e:
        handle_conversion_error(e, "CSV to XLSX")
        raise


def convert_xlsx_to_csv(input_path: Path, output_path: Path):
    """Convert an XLSX file to CSV using pandas."""
    try:
        df = pd.read_excel(str(input_path), engine='openpyxl')
        df.to_csv(str(output_path), index=False)
        logger.info("XLSX to CSV conversion successful: %s", output_path)
        return output_path
    except Exception as e:
        handle_conversion_error(e, "XLSX to CSV")
        raise


# Document conversions - cloud-friendly alternative

def _csv_to_pdf(input_path: Path, output_path: Path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    df = pd.read_csv(str(input_path))
    pdf = SimpleDocTemplate(str(output_path), pagesize=landscape(letter))
    data = [df.columns.tolist()] + df.values.tolist()
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    pdf.build([table])
    logger.info("CSV to PDF conversion successful: %s", output_path)
    return output_path


def _text_to_docx(input_path: Path, output_path: Path):
    """Convert a text file into a DOCX document (one paragraph per non-empty line)."""
    from docx import Document
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    doc = Document()
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(str(output_path))
    logger.info("Text to DOCX conversion successful: %s", output_path)
    return output_path


def _docx_to_text(input_path: Path, output_path: Path):
    """Extract plain text from a DOCX file and write to a text file."""
    from docx import Document
    doc = Document(str(input_path))
    text = '\n\n'.join(
        [para.text for para in doc.paragraphs if para.text.strip()])
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    logger.info("DOCX to Text conversion successful: %s", output_path)
    return output_path


def _text_to_pdf(input_path: Path, output_path: Path):
    """Render a plain text file into a simple PDF document (no complex layout)."""
    from reportlab.lib.pagesizes import letter as PAGE_LETTER
    from reportlab.pdfgen import canvas as pdf_canvas

    c = pdf_canvas.Canvas(str(output_path), pagesize=PAGE_LETTER)
    _, height = PAGE_LETTER
    y = height - 50

    with open(input_path, 'r', encoding='utf-8') as f:
        for line in f:
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(50, y, line.rstrip()[:100])
            y -= 15
    c.save()
    logger.info("Text to PDF conversion successful: %s", output_path)
    return output_path


# Convert with python libs router

def convert_with_python_libs(
        input_path: Path,
        output_path: Path,
        source_ext: str,
        target_format: str):
    try:
        if source_ext == 'csv' and target_format == 'pdf':
            return _csv_to_pdf(input_path, output_path)
        if source_ext in ['txt', 'md', 'html'] and target_format == 'docx':
            return _text_to_docx(input_path, output_path)
        if source_ext == 'docx' and target_format in ['txt', 'md']:
            return _docx_to_text(input_path, output_path)
        if source_ext in ['txt', 'md'] and target_format == 'pdf':
            return _text_to_pdf(input_path, output_path)
        raise ValueError(
            f"Conversion from {source_ext} to {target_format} not supported by Python libs.")
    except Exception as e:
        logger.error("Python libs conversion failed: %s", e)
        handle_conversion_error(e, "Python libs conversion")


# Routing helpers and main router

def _is_pdf_to_doc(source, target):
    return source == 'pdf' and target in ['docx', 'doc']


def _is_doc_to_pdf(source, target):
    return source in ['docx', 'doc'] and target == 'pdf'


def _is_image_to_pdf(source, target):
    return source in ['jpg', 'jpeg', 'png'] and target == 'pdf'


def _is_image_conversion(source, target):
    return source in [
        'jpg',
        'jpeg',
        'png'] and target in [
        'jpg',
        'jpeg',
        'png']


def _is_csv_xlsx_conversion(source, target):
    return (source == 'csv' and target == 'xlsx') or (
        source == 'xlsx' and target == 'csv')


def _is_csv_to_pdf(source, target):
    return source == 'csv' and target == 'pdf'


def convert_file(
        input_path: Path,
        output_path: Path,
        source_ext: str,
        target_ext: str):
    """Dispatch conversions using a small predicate->function mapping to reduce cyclomatic complexity."""
    source = source_ext.lower()
    target = target_ext.lower()

    # pylint: disable=unnecessary-lambda
    dispatch = [
        (lambda s, t: _is_pdf_to_doc(s, t), convert_pdf_to_docx),
        (lambda s, t: _is_doc_to_pdf(s, t), convert_docx_to_pdf),
        (lambda s, t: _is_image_to_pdf(s, t), convert_image_to_pdf),
        (lambda s, t: _is_image_conversion(s, t),
         lambda ip, op: convert_image_to_image(ip, op, target)),
        (lambda s, t: s == 'csv' and t == 'xlsx', convert_csv_to_xlsx),
        (lambda s, t: s == 'xlsx' and t == 'csv', convert_xlsx_to_csv),
        (lambda s, t: _is_csv_to_pdf(s, t),
         lambda ip, op: convert_with_python_libs(ip, op, source, target)),
        (lambda s, t: s in ['txt', 'md', 'html'] and t == 'docx',
         lambda ip, op: convert_with_python_libs(ip, op, source, target)),
        (lambda s, t: s == 'docx' and t in ['txt', 'md'],
         lambda ip, op: convert_with_python_libs(ip, op, source, target)),
        (lambda s, t: s in ['txt', 'md'] and t == 'pdf',
         lambda ip, op: convert_with_python_libs(ip, op, source, target)),
    ]

    for pred, fn in dispatch:
        if pred(source, target):
            return fn(input_path, output_path)

    raise ValueError(
        f"Conversion from {source} to {target} is not supported. Supported: "
        "PDF↔DOCX, Image↔PDF, CSV↔XLSX, CSV→PDF, TXT/MD↔DOCX, TXT→PDF")
