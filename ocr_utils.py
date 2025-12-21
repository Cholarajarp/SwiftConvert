"""Utilities for extracting text from files for OCR and creating simple documents.

This module keeps imports local for heavy optional dependencies (PDF/image parsing)
and provides small helpers used by the Flask app.
"""

# Allow imports inside functions for optional heavy libraries
# pylint: disable=import-outside-toplevel

from pathlib import Path
import logging

from ml_features import extract_text_ocr

logger = logging.getLogger(__name__)


def _create_document_from_text(text: str, output_path: Path, format_type: str):
    """Create document from text in specified format"""
    text = str(text) if text else ''

    if format_type == 'docx':
        from docx import Document
        doc = Document()
        doc.add_paragraph(text)
        doc.save(str(output_path))
        return

    if format_type == 'txt':
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return

    if format_type == 'pdf':
        # Lightweight PDF renderer using ReportLab so OCR->PDF is supported without
        # requiring the converters pipeline.
        from reportlab.lib.pagesizes import letter as PAGE_LETTER
        from reportlab.pdfgen import canvas as pdf_canvas

        c = pdf_canvas.Canvas(str(output_path), pagesize=PAGE_LETTER)
        _, height = PAGE_LETTER
        y = height - 50
        max_width = 80

        for line in str(text).splitlines():
            parts = [line[i:i+max_width] for i in range(0, len(line), max_width)] or ['']
            for part in parts:
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, part)
                y -= 15
        c.save()
        return

    raise ValueError(f"Unsupported format: {format_type}")


def _process_pdf_page(tmp_img: Path):
    """Run OCR on a single page image and return (text, confidence, words)."""
    res = extract_text_ocr(tmp_img, engine='easyocr')
    text = res.get('text', '')
    conf = float(res.get('confidence', 0.0) or 0.0)
    words = len(text.split()) if text else 0
    return text, conf, words


def _process_pdf_for_ocr(input_path: Path, dpi: int = 300):
    """Process PDF pages, run OCR on each page image, and summarize results."""
    import fitz
    logger.info("Processing PDF for OCR: %s (dpi=%s)", input_path, dpi)

    page_texts = []
    total_conf = 0.0
    conf_count = 0
    per_page = []

    with fitz.open(str(input_path)) as pdf_doc:
        for i, page in enumerate(pdf_doc):
            pix = page.get_pixmap(dpi=dpi)
            tmp_img = Path(str(input_path.parent)) / f"{input_path.stem}_page_{i}.png"
            pix.save(str(tmp_img))
            try:
                text, conf, words = _process_pdf_page(tmp_img)
                page_texts.append(text)
                total_conf += conf
                conf_count += 1
                per_page.append({'page': i, 'confidence': round(conf, 3), 'words': words})
            finally:
                try:
                    tmp_img.unlink()
                except OSError as e:
                    logger.debug("Could not remove temp image %s: %s", tmp_img, e)

    combined = ' '.join([t for t in page_texts if t])
    avg_conf = (total_conf / conf_count) if conf_count else 0.0
    return combined, avg_conf, per_page


def _extract_text_for_ocr(input_path: Path, source_ext: str, dpi: int = 300):
    """Extract text from various input types for OCR endpoint.

    Returns (combined_text, ocr_result_dict)
    """
    ext = source_ext.lower()
    ocr_result = {'text': '', 'confidence': 0.0, 'word_count': 0}

    if ext == 'pdf':
        combined, avg_conf, per_page = _process_pdf_for_ocr(input_path, dpi)
        ocr_result = {
            'text': combined,
            'confidence': round(avg_conf, 3),
            'word_count': len(combined.split()),
            'per_page': per_page,
        }
        return combined, ocr_result

    if ext in ['docx', 'doc']:
        logger.info("Extracting text from DOCX: %s", input_path)
        from docx import Document as DocxReader
        docx_obj = DocxReader(str(input_path))
        paragraphs = [p.text for p in docx_obj.paragraphs if p.text]
        combined = '\n'.join(paragraphs)
        ocr_result = {
            'text': combined,
            'confidence': 1.0,
            'word_count': len(combined.split())
        }
        return combined, ocr_result

    if ext in ['txt', 'md', 'html']:
        logger.info("Reading text file: %s", input_path)
        text = input_path.read_text(encoding='utf-8', errors='ignore')
        if ext == 'html':
            import re
            text = re.sub(r'<[^>]+>', '', text)
        ocr_result = {
            'text': text,
            'confidence': 1.0,
            'word_count': len(text.split())
        }
        return text, ocr_result

    logger.info("Running image OCR on: %s", input_path)
    res = extract_text_ocr(input_path, engine='easyocr')
    text = res.get('text', '')
    conf = float(res.get('confidence', 0.0) or 0.0)
    ocr_result = {
        'text': text, 'confidence': round(
            conf, 3), 'word_count': len(
            text.split())}
    return text, ocr_result
