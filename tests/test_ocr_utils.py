from ocr_utils import _extract_text_for_ocr, _create_document_from_text
from docx import Document


def test_extract_text_docx(tmp_path):
    doc = Document()
    doc.add_paragraph('Hello World')
    doc.add_paragraph('Second line')
    p = tmp_path / 'sample.docx'
    doc.save(str(p))

    text, info = _extract_text_for_ocr(p, 'docx')
    assert 'Hello World' in text
    assert info['confidence'] == 1.0
    assert info['word_count'] >= 2


def test_extract_text_txt_md_html(tmp_path):
    txt = tmp_path / 'sample.txt'
    txt.write_text('Line one\nLine two', encoding='utf-8')

    text, info = _extract_text_for_ocr(txt, 'txt')
    assert 'Line one' in text
    assert info['confidence'] == 1.0

    md = tmp_path / 'sample.md'
    md.write_text('# Title\nContent here', encoding='utf-8')
    text_md, info_md = _extract_text_for_ocr(md, 'md')
    assert 'Title' in text_md or 'Content' in text_md

    html = tmp_path / 'sample.html'
    html.write_text('<h1>Hi</h1><p>Paragraph</p>', encoding='utf-8')
    text_html, info_html = _extract_text_for_ocr(html, 'html')
    assert 'Paragraph' in text_html


# Image and PDF flows require heavy dependencies; we test docx/txt/md/html which are lightweight.

def test_create_document_pdf(tmp_path):
    out_pdf = tmp_path / 'out.pdf'
    text = 'Hello World\nThis is a test PDF generated from text.'
    _create_document_from_text(text, out_pdf, 'pdf')
    assert out_pdf.exists(), 'PDF file should be created'
    assert out_pdf.stat().st_size > 0
